from __future__ import annotations

from datetime import datetime
import json
import logging
from pathlib import Path
import sys
from typing import Iterator

from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response


REPO_ROOT = Path(__file__).resolve().parents[2]
PYTHON_SRC = REPO_ROOT / "packages" / "python" / "src"
if str(PYTHON_SRC) not in sys.path:
    sys.path.insert(0, str(PYTHON_SRC))

from law_rag_core.auth import AuthService
from law_rag_core.chat import ChatService
from law_rag_core.db import ArtifactNotReadyError, get_state_db_connection
from law_rag_core.ingest import IngestService
from law_rag_core.logging import configure_logging
from law_rag_core.precedent import PrecedentService
from law_rag_core.repository import Repository
from law_rag_core.settings import get_settings
from law_rag_core.types import (
    AuthResult,
    ChatRequest,
    FeedbackCreate,
    IngestRunRequest,
    PrecedentDetailResponse,
    PrecedentSearchResponse,
    PrecedentSource,
    SavedAnswerCreate,
)


configure_logging()
logger = logging.getLogger(__name__)
settings = get_settings()
repository = Repository()
chat_service = ChatService()
ingest_service = IngestService()
precedent_service = PrecedentService()
auth_service = AuthService(repository=repository)

app = FastAPI(title="KR Law RAG MVP", version="0.1.0")


# ── Security Headers Middleware (helmet 동급) ─────────────────
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """브라우저 보안 헤더를 모든 응답에 자동 추가. Node.js의 helmet과 동일한 역할."""

    async def dispatch(self, request: Request, call_next) -> Response:
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
        if request.url.scheme == "https":
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response


app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_allow_origins,
    allow_credentials=settings.cors_allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
)


def require_admin(x_admin_api_key: str | None = Header(default=None)) -> None:
    """관리자 인증 — X-Admin-Api-Key 헤더 검증. 불일치 시 401 반환."""
    if not settings.admin_api_key or x_admin_api_key != settings.admin_api_key:
        raise HTTPException(status_code=401, detail="Invalid admin API key.")


def get_optional_user(authorization: str | None = Header(default=None)) -> dict | None:
    """선택적 사용자 인증 — Authorization Bearer 토큰에서 사용자 정보 추출. 토큰 없거나 무효하면 None 반환."""
    return auth_service.get_current_user(authorization)


def _client_ip(request: Request) -> str:
    """요청에서 클라이언트 IP를 추출한다. X-Forwarded-For → client.host 순서."""
    return request.headers.get("x-forwarded-for", "").split(",")[0].strip() or (
        request.client.host if request.client else "unknown"
    )


def _record_ops_event(
    event_type: str,
    *,
    channel: str = "web",
    user_id: str | None = None,
    session_id: str | None = None,
    metadata: dict | None = None,
) -> None:
    """운영 대시보드용 이벤트를 저장한다. 실패해도 본 요청은 계속 처리한다."""
    try:
        repository.record_request_event(
            event_type,
            channel=channel,
            user_id=user_id,
            session_id=session_id,
            metadata=metadata or {},
        )
    except Exception as exc:
        logger.warning("Failed to record ops event %s: %s", event_type, exc)


def _usage_date_key() -> str:
    return datetime.now(settings.timezone_info).strftime("%Y-%m-%d")


@app.on_event("startup")
def ensure_state_schema() -> None:
    """서버 시작 시 state DB 스키마를 보장한다."""
    try:
        with get_state_db_connection():
            pass
    except Exception as exc:
        logger.warning("Failed to ensure state DB schema on startup: %s", exc)


@app.get("/v1/health")
def health() -> dict[str, str]:
    """헬스체크 — 서비스 상태 + 법령 DB/벡터 아티팩트 준비 여부 확인.
    반환: {"status", "env", "law_artifact", "vector_artifact"}
    """
    return {
        "status": "ok",
        "env": settings.app_env,
        "law_artifact": "ready" if settings.law_db_path.exists() else "missing",
        "vector_artifact": "ready" if settings.vector_matrix_path.exists() else "missing",
    }


# ── Auth endpoints ───────────────────────────────────────────


class KakaoLoginRequest(BaseModel):
    code: str


@app.post("/v1/auth/kakao", response_model=AuthResult)
def kakao_login(payload: KakaoLoginRequest):
    """카카오 OAuth 로그인 — 인가 코드(code)를 받아 JWT access_token + 사용자 정보 반환.
    수신: KakaoLoginRequest(code: str)
    반환: AuthResult(access_token, user_id, nickname 등)
    """
    try:
        result = auth_service.login_or_register(payload.code)
        _record_ops_event(
            "auth_login",
            channel="web",
            user_id=result.user_id,
            metadata={"is_new": result.is_new},
        )
        return result
    except RuntimeError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/v1/auth/me")
def auth_me(user: dict | None = Depends(get_optional_user)):
    """현재 로그인 사용자 정보 반환 — JWT 토큰에서 user_id, nickname 추출.
    인증 실패 시 401 반환.
    """
    if user is None:
        raise HTTPException(status_code=401, detail="Invalid or missing token.")
    return {"user_id": user["sub"], "nickname": user.get("nickname")}


@app.get("/v1/usage")
def get_usage(
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """일일 사용량 조회 — 로그인 사용자는 user_id 기준, 비로그인은 session_id 기준.
    반환: {"used": 현재 사용 횟수, "limit": 일일 한도, "logged_in": bool}
    """
    if user:
        count = repository.get_usage_count(user_id=user["sub"])
        # Fallback: if user has 0 but session has counts (auth was lost during chat)
        if count == 0 and x_session_id:
            session_count = repository.get_usage_count(session_id=x_session_id)
            count = max(count, session_count)
        return {"used": count, "limit": USER_LIMIT, "logged_in": True}
    session_id = x_session_id or "anonymous"
    count = repository.get_usage_count(session_id=session_id)
    return {"used": count, "limit": GUEST_LIMIT, "logged_in": False}


# ── Chat ─────────────────────────────────────────────────────


GUEST_LIMIT = 3
USER_LIMIT = 30  # per-day (changed from permanent 10 to daily 30 for data accumulation)
IP_DAILY_LIMIT = 5  # IP-based hard cap even with session rotation


def _check_daily_limit() -> None:
    """전체 시스템 일일 한도 검증 (비용 보호) — 오늘 전체 질문 수가 daily_question_limit 초과 시 503 반환."""
    import sqlite3
    today = _usage_date_key()
    try:
        conn = sqlite3.connect(settings.state_db_path)
        c = conn.cursor()
        c.execute("SELECT SUM(count) FROM usage_counts WHERE date = ? AND session_id NOT LIKE 'ip:%'", (today,))
        total = c.fetchone()[0] or 0
        conn.close()
        if total >= settings.daily_question_limit:
            raise HTTPException(status_code=503, detail="일일 서비스 한도에 도달했습니다. 내일 다시 이용해 주세요.")
    except HTTPException:
        raise
    except Exception:
        pass


def _format_review_result(result: dict, extra: dict | None = None) -> dict:
    """review_document/review_documents_orchestrated 결과를 API 응답 형식으로 변환."""
    overview = result.get("document_overview") or {}
    resp = {
        "documentOverview": overview,
        "documentType": overview.get("nature") or result.get("document_type", ""),
        "summary": result.get("summary", ""),
        "institutionalFrame": result.get("institutional_frame", ""),
        "institutionalAxes": result.get("axes", []),
        "reviewAxes": result.get("review_axes", []),
        "observations": result.get("observations", []),
        "clauseReviews": result.get("clause_reviews", []),
        "gaps": result.get("gaps", []),
        "missingClauses": result.get("missing_clauses", []),
        "externalConsiderations": result.get("external_considerations", []),
        "riskScenarios": result.get("risk_scenarios", []),
        "overallRisk": result.get("overall_risk", "n/a"),
        "keyPoints": result.get("key_points", []),
        "actionItems": result.get("action_items", []),
        "rejectedCitations": result.get("rejected_citations", []),
        "judgment": result.get("judgment") or None,
        "disclaimer": result.get("disclaimer", ""),
    }
    if extra:
        resp.update(extra)
    return resp


def _record_usage(user: dict | None, x_session_id: str | None, request: Request) -> None:
    """응답 성공 후 사용량 증가 기록."""
    if user:
        repository.increment_usage(user_id=user["sub"])
    else:
        session_id = x_session_id or "anonymous"
        repository.increment_usage(session_id=session_id)
        repository.increment_usage(session_id=f"ip:{_client_ip(request)}")


@app.post("/v1/chat")
def chat(
    payload: ChatRequest,
    request: Request,
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """채팅 엔드포인트 — 사용량 검증 후 ChatService.answer() 호출.
    수신: ChatRequest(question, stream 등)
    반환: stream=True → SSE(text/event-stream), stream=False → JSON(AnswerResult)
    한도 초과 시 429, 아티팩트 미준비 시 503 반환.
    """
    # TODO: 사용량 제한 임시 비활성화 — 평가 완료 후 복원할 것
    # _check_daily_limit()
    #
    # if user:
    #     user_id = user["sub"]
    #     current = repository.get_usage_count(user_id=user_id)
    #     if current >= USER_LIMIT:
    #         raise HTTPException(status_code=429, detail="LIMIT_EXCEEDED")
    # else:
    #     session_id = x_session_id or "anonymous"
    #     session_count = repository.get_usage_count(session_id=session_id)
    #     if session_count >= GUEST_LIMIT:
    #         raise HTTPException(status_code=429, detail="LOGIN_REQUIRED")
    #     client_ip = request.headers.get("x-forwarded-for", "").split(",")[0].strip() or request.client.host if request.client else "unknown"
    #     ip_key = f"ip:{client_ip}"
    #     ip_count = repository.get_usage_count(session_id=ip_key)
    #     if ip_count >= IP_DAILY_LIMIT:
    #         raise HTTPException(status_code=429, detail="LOGIN_REQUIRED")

    try:
        answer = chat_service.answer(payload)
    except ArtifactNotReadyError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover - surfaced to client
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    _record_usage(user, x_session_id, request)

    # 통합 히스토리 저장 (chat 유형)
    try:
        repository.save_chat_history(
            user_id=user["sub"] if user else "anonymous",
            question=payload.question,
            response=answer.model_dump(mode="json"),
            session_id=x_session_id,
            request_type="chat",
        )
    except Exception:
        pass

    _record_ops_event(
        "chat_answer",
        channel=payload.channel,
        user_id=user["sub"] if user else None,
        session_id=None if user else (x_session_id or "anonymous"),
        metadata={
            "stream": payload.stream,
            "grounded": answer.grounded,
            "citations": len(answer.citations),
            "precedents": len(answer.precedents),
        },
    )

    if not payload.stream:
        return answer

    # SSE 스트리밍: meta → token(청크별) → citations → done 순서로 전송
    def event_stream() -> Iterator[str]:
        yield f"event: meta\ndata: {json.dumps({'summary': answer.summary, 'grounded': answer.grounded}, ensure_ascii=False)}\n\n"
        for token in _chunk_text(answer.answer):
            yield f"event: token\ndata: {json.dumps({'text': token}, ensure_ascii=False)}\n\n"
        yield f"event: citations\ndata: {json.dumps([citation.model_dump(mode='json') for citation in answer.citations], ensure_ascii=False)}\n\n"
        yield f"event: done\ndata: {json.dumps(answer.model_dump(mode='json'), ensure_ascii=False)}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ── Ouroboros Deep-Dive (AI-driven follow-up questions) ──────


class DeepDiveAnswerTurn(BaseModel):
    """한 질문에 대한 사용자 답변 (single/multi/yes_no/free_text 공통 래퍼)."""
    question_id: str
    question_text: str
    input_type: str = "free_text"  # "single_choice" | "multi_choice" | "yes_no" | "free_text"
    selected_options: list[str] = []  # multi_choice·single_choice·yes_no 선택 결과
    free_text: str | None = None  # allow_free_text 로 입력된 자유 텍스트


class DeepDiveRequest(BaseModel):
    """Ouroboros pattern: send message + conversation history + 이번 턴 답변 묶음."""
    message: str
    conversation: list[dict[str, str]] = []  # [{"role": "user"/"assistant", "content": "..."}]
    answers: list[DeepDiveAnswerTurn] = []  # 이번 턴에 사용자가 응답한 질문들
    force_analyze: bool = False  # true 면 follow-up 건너뛰고 즉시 분석
    thread_id: str | None = None  # 대화 스레드 ID (프론트에서 전달, 없으면 서버에서 생성)


# 하드캡: 수집 턴이 이 수를 넘으면 서버가 강제로 ready=true 로 전환하여
# 결과 도출 파이프라인으로 진입시킨다. conversation 은 user+assistant 쌍이므로
# 6 = 3턴(사용자가 답변한 횟수) 기준.
DEEP_DIVE_MAX_CONVERSATION_MESSAGES = 6


def _render_answers_block(answers: list[DeepDiveAnswerTurn]) -> str:
    """이번 턴 answers 를 대화 턴으로 삽입할 한국어 블록으로 렌더링한다."""
    if not answers:
        return ""
    lines: list[str] = []
    for a in answers:
        picked = list(a.selected_options or [])
        if a.free_text:
            picked.append(f"(직접입력) {a.free_text}")
        value = ", ".join(picked) if picked else "(응답 없음)"
        lines.append(f"- {a.question_text} → {value}")
    return "\n".join(lines)


def _compose_analysis_document(
    *,
    original_message: str,
    conversation: list[dict[str, str]],
    final_answers: list[DeepDiveAnswerTurn],
    gathered_facts: dict[str, str],
) -> str:
    """수집된 사실관계·대화 이력·이번 답변을 '심층 검토용 합성 문서'로 조립.

    이 문서는 C(`/v1/analyze-text`)의 review_document 엔진 입력으로 그대로 쓰인다.
    """
    sections: list[str] = []
    sections.append("[사용자 원 질문]\n" + original_message.strip())

    if gathered_facts:
        fact_lines = [f"- {k}: {v}" for k, v in gathered_facts.items() if v]
        if fact_lines:
            sections.append("[수집된 사실관계]\n" + "\n".join(fact_lines))

    answers_block = _render_answers_block(final_answers)
    if answers_block:
        sections.append("[이번 턴 응답]\n" + answers_block)

    if conversation:
        conv_lines: list[str] = []
        for m in conversation:
            role = "사용자" if m.get("role") == "user" else "AI"
            content = (m.get("content") or "").strip()
            if content:
                conv_lines.append(f"{role}: {content}")
        if conv_lines:
            sections.append("[대화 이력]\n" + "\n".join(conv_lines))

    return "\n\n".join(sections)


@app.post("/v1/deep-dive")
def deep_dive(
    payload: DeepDiveRequest,
    request: Request,
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """대화형 심층 분석 (Ouroboros 패턴) — AI가 추가 질문으로 사실관계 수집 후 구조화된 법률 분석 제공.
    수신: DeepDiveRequest(message, conversation[])
    반환: 추가 질문 필요 시 {"type": "follow_up", "questions": [...]},
          분석 준비 완료 시 {"type": "analysis", "summary", "answer", "favorableFacts", ...}
    """
    # TODO: 사용량 제한 임시 비활성화 — 평가 완료 후 복원할 것
    # _check_daily_limit()
    #
    # if user:
    #     user_id = user["sub"]
    #     current = repository.get_usage_count(user_id=user_id)
    #     if current >= USER_LIMIT:
    #         raise HTTPException(status_code=429, detail="LIMIT_EXCEEDED")
    # elif not payload.conversation:
    #     session_id = x_session_id or "anonymous"
    #     session_count = repository.get_usage_count(session_id=session_id)
    #     if session_count >= GUEST_LIMIT:
    #         raise HTTPException(status_code=429, detail="LOGIN_REQUIRED")

    try:
        from law_rag_core.ai import GeminiService
        from law_rag_core.retrieval import RetrievalService
        gemini = GeminiService()

        # 스레드 ID: 프론트에서 전달하거나 첫 턴이면 서버에서 생성
        import uuid as _uuid
        thread_id = payload.thread_id or str(_uuid.uuid4())
        turn_index = len(payload.conversation) // 2  # user+assistant 쌍 기준

        # 이번 턴 answers 를 대화 이력에 사용자 발화로 추가 (follow-up LLM 에 전달하기 위함).
        # 프론트는 "질문 묶음에 답한 뒤 다음 턴" 으로 호출하므로, 서버가 answers 를
        # 대화 텍스트로 렌더해 conversation 에 append 한 뒤 LLM 에 넘긴다.
        conversation_for_llm = list(payload.conversation)
        answers_text = _render_answers_block(payload.answers)
        if answers_text:
            conversation_for_llm.append({
                "role": "user",
                "content": f"(이전 턴 질문에 대한 응답)\n{answers_text}",
            })

        # ── 턴 하드캡 / 사용자 강제종료 판정 ──
        # 1) force_analyze: 사용자가 '이제 분석' 버튼을 누른 경우
        # 2) DEEP_DIVE_MAX_CONVERSATION_MESSAGES 초과: 서버 보호용 하드캡
        force_ready = (
            payload.force_analyze
            or len(conversation_for_llm) >= DEEP_DIVE_MAX_CONVERSATION_MESSAGES
        )

        # Step 1: follow-up 판단 (force_ready 면 LLM 스킵)
        if force_ready:
            result = {
                "ready": True,
                "gathered_facts": {},
                "domain": "일반",
                "analysis_question": payload.message,
            }
        else:
            result = gemini.generate_follow_up_questions(
                question=payload.message,
                conversation_history=conversation_for_llm,
            )

        if not result.get("ready", False):
            # 아직 정보 부족 — 질문 묶음 반환
            questions_raw = result.get("questions", []) or []
            # 스키마 정규화: 구(type) → 신(input_type) 마이그레이션 호환
            questions_norm: list[dict] = []
            for q in questions_raw:
                if not isinstance(q, dict):
                    continue
                input_type = q.get("input_type") or q.get("type") or "free_text"
                if input_type == "choice":
                    input_type = "single_choice"
                questions_norm.append({
                    "id": q.get("id") or f"q{len(questions_norm)+1}",
                    "question": q.get("question", ""),
                    "input_type": input_type,
                    "options": q.get("options") or [],
                    "allow_free_text": bool(q.get("allow_free_text", False)),
                    "legal_reason": q.get("legal_reason", ""),
                })

            follow_up_resp = {
                "type": "follow_up",
                "message": result.get("message", ""),
                "questions": questions_norm,
                "gatheredFacts": result.get("gathered_facts", {}),
                "thread_id": thread_id,
                "turn_index": turn_index,
                "max_turns": DEEP_DIVE_MAX_CONVERSATION_MESSAGES // 2,
            }
            try:
                repository.save_chat_history(
                    user_id=user["sub"] if user else "anonymous",
                    question=payload.message,
                    response=follow_up_resp,
                    session_id=x_session_id,
                    thread_id=thread_id,
                    turn_index=turn_index,
                    request_type="deep_dive_follow_up",
                )
            except Exception:
                pass
            _record_ops_event(
                "deep_dive_follow_up",
                channel="web",
                user_id=user["sub"] if user else None,
                session_id=None if user else (x_session_id or "anonymous"),
                metadata={
                    "question_count": len(questions_norm),
                    "conversation_turns": len(conversation_for_llm),
                    "thread_id": thread_id,
                },
            )
            return follow_up_resp

        # ── Step 2: ready=true — 합성 문서 조립 후 C 엔진(review_document) 재사용 ──
        gathered_facts = result.get("gathered_facts", {}) or {}
        domain = result.get("domain", "일반")

        composed_text = _compose_analysis_document(
            original_message=payload.message,
            conversation=conversation_for_llm,
            final_answers=payload.answers,
            gathered_facts=gathered_facts,
        )

        # analysis_question: retrieve / review_document 에 쓰이는 분석 질문 문자열.
        # LLM 이 작성한 요약(analysis_question) 이 있으면 우선 사용, 없으면 원 질문.
        analysis_question = (
            result.get("analysis_question")
            or payload.message
            or "수집된 사실관계에 대한 법적 검토와 판단을 수행해 주세요."
        )

        retrieval = RetrievalService()
        # 합성 문서 + 분석 질문을 쿼리로 사용해 recall 강화
        retrieve_query = f"{analysis_question}\n\n{composed_text}"
        chunks = retrieval.retrieve(retrieve_query)

        # C 엔진: 문서 검토 파이프라인 (Multi-Chain). 대화형 수집 결과를 동일 품질로 검토.
        review_result = gemini.review_document(
            document_text=composed_text,
            question=analysis_question,
            chunks=chunks,
        )
        response = _format_review_result(review_result, extra={
            "type": "analysis",
            "thread_id": thread_id,
            "domain": domain,
            "gatheredFacts": gathered_facts,
            "turn_index": turn_index,
            "force_analyzed": force_ready,
        })

        # 사용량 증가
        if user:
            repository.increment_usage(user_id=user["sub"])
        else:
            session_id = x_session_id or "anonymous"
            repository.increment_usage(session_id=session_id)
            client_ip = _client_ip(request)
            repository.increment_usage(session_id=f"ip:{client_ip}")

        try:
            repository.save_chat_history(
                user_id=user["sub"] if user else "anonymous",
                question=analysis_question,
                response=response,
                session_id=x_session_id,
                thread_id=thread_id,
                turn_index=turn_index,
                request_type="deep_dive_analysis",
            )
        except Exception:
            pass
        _record_ops_event(
            "deep_dive_analysis",
            channel="web",
            user_id=user["sub"] if user else None,
            session_id=None if user else (x_session_id or "anonymous"),
            metadata={
                "observations": len(response.get("observations", []) or []),
                "clause_reviews": len(response.get("clauseReviews", []) or []),
                "judgment": bool(response.get("judgment")),
                "force_analyzed": force_ready,
                "thread_id": thread_id,
            },
        )

        return response

    except ArtifactNotReadyError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


# ── File Analysis ─────────────────────────────────────────────


MAX_UPLOAD_SIZE = 20 * 1024 * 1024  # 20MB


def _extract_text_from_upload(file: UploadFile) -> str:
    """업로드 파일에서 텍스트 추출 (txt/md/csv/json/docx/hwp/hwpx/pdf/xlsx/이미지 지원).
    수신: UploadFile 객체
    반환: 추출된 텍스트 문자열
    20MB 초과 시 413, 미지원 형식 시 400, 라이브러리 미설치 시 500 반환.
    """
    content = file.file.read()
    if len(content) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=413, detail=f"파일 크기가 {MAX_UPLOAD_SIZE // 1024 // 1024}MB를 초과합니다.")
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()

    # 텍스트 기반 파일은 UTF-8로 직접 디코딩
    if suffix in (".txt", ".md", ".csv", ".json"):
        return content.decode("utf-8", errors="ignore")

    # DOCX: python-docx로 본문 + 표 텍스트 추출
    if suffix == ".docx":
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return "\n".join(paragraphs)
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX 처리 라이브러리가 설치되지 않았습니다.")
        except Exception as e:
            if "password" in str(e).lower() or "encrypt" in str(e).lower():
                raise HTTPException(status_code=400, detail="비밀번호가 설정된 파일입니다. 비밀번호를 해제한 후 다시 업로드해 주세요.")
            raise HTTPException(status_code=400, detail=f"DOCX 처리 실패: {str(e)}")

    if suffix == ".hwpx":
        # hwpx는 ZIP/XML 기반
        try:
            import zipfile
            import io
            from xml.etree import ElementTree
            zf = zipfile.ZipFile(io.BytesIO(content))
            text_parts = []
            for name in sorted(zf.namelist()):
                if name.startswith("Contents/") and name.endswith(".xml"):
                    xml_data = zf.read(name)
                    root = ElementTree.fromstring(xml_data)
                    for elem in root.iter():
                        if elem.text and elem.text.strip():
                            text_parts.append(elem.text.strip())
            zf.close()
            text = "\n".join(text_parts)
            if text.strip():
                return text
            raise HTTPException(status_code=400, detail="HWPX 파일에서 텍스트를 추출할 수 없습니다.")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=400, detail="HWPX 파일 처리에 실패했습니다. PDF 또는 DOCX로 변환 후 업로드해 주세요.")

    # HWP (한글 바이너리): olefile로 OLE 스트림 읽기 → zlib 해제 → UTF-16LE 텍스트 추출
    if suffix == ".hwp":
        try:
            import olefile
            import io
            import zlib
            if not olefile.isOleFile(io.BytesIO(content)):
                raise HTTPException(status_code=400, detail="HWP 파일 형식이 올바르지 않습니다.")
            ole = olefile.OleFileIO(io.BytesIO(content))
            if ole.exists("BodyText/Section0"):
                text_parts = []
                for i in range(256):
                    stream_name = f"BodyText/Section{i}"
                    if not ole.exists(stream_name):
                        break
                    data = ole.openstream(stream_name).read()
                    try:
                        data = zlib.decompress(data, -15)
                    except zlib.error:
                        pass
                    # HWP 바이너리에서 텍스트 추출 (UTF-16LE)
                    text = ""
                    pos = 0
                    while pos < len(data):
                        if pos + 1 < len(data):
                            ch = int.from_bytes(data[pos:pos+2], 'little')
                            if 32 <= ch < 0xFFFF and ch not in (0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31) or ch in (10, 13):
                                text += chr(ch)
                        pos += 2
                    text_parts.append(text)
                result_text = "\n".join(text_parts)
                # 서로게이트 문자 및 제어 문자 제거
                result_text = result_text.encode("utf-8", errors="ignore").decode("utf-8")
                result_text = "".join(ch for ch in result_text if ord(ch) >= 32 or ch in "\n\r\t")
                ole.close()
                if result_text.strip():
                    return result_text
            ole.close()
            raise HTTPException(status_code=400, detail="HWP 파일에서 텍스트를 추출할 수 없습니다. PDF 또는 DOCX로 변환 후 업로드해 주세요.")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=400, detail="HWP 파일 처리에 실패했습니다. PDF 또는 DOCX로 변환 후 업로드해 주세요.")

    if suffix == ".doc":
        # .doc (구 Word 형식) — textract 또는 antiword로 처리, 없으면 Gemini로 추출
        try:
            import subprocess
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            # antiword 시도
            try:
                result = subprocess.run(
                    ["antiword", tmp_path], capture_output=True, text=True, timeout=30,
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except FileNotFoundError:
                pass
            # libreoffice --headless 변환 시도
            try:
                import os
                out_dir = tempfile.mkdtemp()
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", out_dir, tmp_path],
                    capture_output=True, timeout=30,
                )
                txt_path = os.path.join(out_dir, os.path.splitext(os.path.basename(tmp_path))[0] + ".txt")
                if os.path.exists(txt_path):
                    with open(txt_path, encoding="utf-8", errors="ignore") as f:
                        return f.read()
            except FileNotFoundError:
                pass
            # 최후 수단: Gemini로 바이너리에서 텍스트 추출 시도
            from law_rag_core.ai import GeminiService
            from google.genai import types
            gemini = GeminiService()
            if gemini.client:
                response = gemini.client.models.generate_content(
                    model=gemini.settings.gemini_model,
                    contents=[
                        types.Part.from_bytes(data=content, mime_type="application/msword"),
                        "이 Word 문서(.doc)의 텍스트를 모두 추출하라. 텍스트만 반환.",
                    ],
                    config=types.GenerateContentConfig(temperature=0.0),
                )
                if response.text:
                    return response.text
            raise HTTPException(status_code=400, detail=".doc 파일을 처리할 수 없습니다. .docx 형식으로 변환 후 다시 업로드해 주세요.")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f".doc 파일 처리 실패: {str(e)}. .docx 형식으로 변환 후 다시 업로드해 주세요.")

    # PDF: pdfplumber로 텍스트+표 추출, 스캔 PDF는 Gemini Vision OCR 시도
    if suffix == ".pdf":
        try:
            import pdfplumber
            import io
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages[:50]:
                    page_text = page.extract_text() or ""
                    # 표도 추출
                    for table in (page.extract_tables() or []):
                        for row in table:
                            cells = [c or "" for c in row]
                            page_text += "\n" + " | ".join(cells)
                    text_parts.append(page_text)
            text = "\n".join(text_parts)
            if not text.strip():
                # 스캔 PDF (이미지만) → Gemini Vision OCR 시도
                from law_rag_core.ai import GeminiService
                from google.genai import types
                gemini = GeminiService()
                if gemini.client:
                    response = gemini.client.models.generate_content(
                        model=gemini.settings.gemini_model,
                        contents=[
                            types.Part.from_bytes(data=content, mime_type="application/pdf"),
                            "이 PDF의 모든 텍스트를 추출하라. 표는 표 형태로, 스캔된 이미지의 텍스트도 모두 포함. 텍스트만 반환.",
                        ],
                        config=types.GenerateContentConfig(temperature=0.0),
                    )
                    if response.text:
                        return response.text
                raise HTTPException(status_code=400, detail="PDF에서 텍스트를 추출할 수 없습니다. 텍스트가 포함된 PDF를 업로드해 주세요.")
            return text
        except HTTPException:
            raise
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF 처리 라이브러리가 설치되지 않았습니다.")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF 처리 실패: {str(e)}")

    # 엑셀: openpyxl로 시트별 셀 데이터 추출
    if suffix in (".xlsx", ".xls"):
        try:
            import openpyxl
            import io
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text_parts.append(f"[시트: {sheet}]")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c for c in cells):
                        text_parts.append(" | ".join(cells))
            return "\n".join(text_parts)
        except ImportError:
            raise HTTPException(status_code=400, detail="엑셀 파일은 현재 지원하지 않습니다. PDF 또는 DOCX로 변환 후 업로드해 주세요.")
        except Exception:
            raise HTTPException(status_code=400, detail="엑셀 파일 처리에 실패했습니다.")

    # 이미지: Gemini Vision으로 OCR 텍스트 추출
    if suffix in (".png", ".jpg", ".jpeg", ".webp"):
        from law_rag_core.ai import GeminiService
        from google.genai import types
        gemini = GeminiService()
        if not gemini.client:
            raise HTTPException(status_code=503, detail="AI 서비스가 설정되지 않았습니다.")
        mime = f"image/{suffix[1:]}" if suffix != ".jpg" else "image/jpeg"
        response = gemini.client.models.generate_content(
            model=gemini.settings.gemini_model,
            contents=[
                types.Part.from_bytes(data=content, mime_type=mime),
                "이 이미지의 모든 텍스트를 추출하라. 표는 표 형태로, 도장/서명은 [도장]/[서명]으로 표시. 텍스트만 반환.",
            ],
            config=types.GenerateContentConfig(temperature=0.0),
        )
        return response.text or ""

    supported = "txt, pdf, docx, doc, hwp, hwpx, xlsx, png, jpg, jpeg, webp"
    raise HTTPException(status_code=400, detail=f"지원하지 않는 파일 형식입니다 ({suffix}). 지원 형식: {supported}")


@app.post("/v1/analyze-file")
async def analyze_file(
    files: list[UploadFile] = File(...),
    question: str = Form(default=""),
    previous_summary: str = Form(default=""),
    request: Request = None,
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """파일 업로드 분석 — 최대 5개 파일 텍스트 추출 → RAG 검색 → Gemini 문서 검토 파이프라인.
    수신: files(업로드 파일), question(분석 질문, 선택), previous_summary(이전 검토 결과, 선택)
    반환: {"files", "documentOverview", "summary", "clauseReviews", "gaps", "actionItems", ...}
    """
    # TODO: 사용량 제한 임시 비활성화 — 평가 완료 후 복원할 것
    # _check_daily_limit()

    # 최대 5개 파일
    if len(files) > 5:
        raise HTTPException(status_code=400, detail="파일은 최대 5개까지 첨부할 수 있습니다.")

    try:
        # 각 파일에서 텍스트 추출
        extracted_files: list[dict] = []
        combined_text_parts: list[str] = []

        for file in files:
            file_text = _extract_text_from_upload(file)
            if not file_text.strip():
                continue
            extracted_files.append({
                "filename": file.filename or "unknown",
                "extractedText": file_text[:2000],
                "extractedLength": len(file_text),
            })
            combined_text_parts.append(f"--- {file.filename or 'unknown'} ---\n{file_text}")

        if not extracted_files:
            raise HTTPException(status_code=400, detail="파일에서 텍스트를 추출할 수 없습니다.")

        combined_text = "\n\n".join(combined_text_parts)

        # 질문 구성 (이전 검토 결과가 있으면 맥락으로 포함)
        user_question = question or "이 문서를 법적으로 분석해 주세요."
        if previous_summary.strip():
            user_question = f"{user_question}\n\n[이전 검토 결과 요약]\n{previous_summary[:2000]}"

        # RAG 검색 (질문 기반)
        from law_rag_core.retrieval import RetrievalService
        from law_rag_core.ai import GeminiService
        retrieval = RetrievalService()
        gemini = GeminiService()

        chunks = retrieval.retrieve(user_question if question else combined_text[:200])

        # 문서 검토: 파일 1개면 단일 검토, 여러 개면 오케스트레이션
        if len(extracted_files) == 1:
            result = gemini.review_document(
                document_text=combined_text,
                question=user_question,
                chunks=chunks,
            )
        else:
            doc_list = [
                {"filename": ef["filename"], "text": t}
                for ef, t in zip(extracted_files, combined_text_parts)
            ]
            result = gemini.review_documents_orchestrated(
                documents=doc_list,
                question=user_question,
                chunks=chunks,
            )

        # 결과에 문서 텍스트 해시 포함 (추가 질문 시 연결용)
        import hashlib
        result["document_hash"] = hashlib.md5(combined_text.encode()).hexdigest()[:12]

        overview = result.get("document_overview") or {}

        file_names = [f["name"] for f in extracted_files]
        # 통합 히스토리 저장 (파일 검토 유형)
        try:
            repository.save_chat_history(
                user_id=user["sub"] if user else "anonymous",
                question=question,
                response={"summary": result.get("summary", ""), "overall_risk": result.get("overall_risk", "n/a")},
                session_id=x_session_id,
                request_type="review_file",
                has_attachments=True,
                attachment_names=file_names,
            )
        except Exception:
            pass

        _record_ops_event(
            "review_file",
            channel="web",
            user_id=user["sub"] if user else None,
            session_id=None if user else (x_session_id or "anonymous"),
            metadata={
                "file_count": len(extracted_files),
                "total_extracted_length": sum(f["extractedLength"] for f in extracted_files),
                "overall_risk": result.get("overall_risk", "n/a"),
            },
        )
        return _format_review_result(result, {
            "files": extracted_files,
            "fileCount": len(extracted_files),
            "totalExtractedLength": sum(f["extractedLength"] for f in extracted_files),
        })

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


class AnalyzeTextRequest(BaseModel):
    text: str
    question: str = ""
    previous_summary: str = ""


@app.post("/v1/analyze-text")
def analyze_text(
    payload: AnalyzeTextRequest,
    request: Request,
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """텍스트 직접 분석 — 파일 없이 본문 텍스트(100자 이상)로 심층 검토 파이프라인 실행.
    수신: AnalyzeTextRequest(text, question, previous_summary)
    반환: analyze-file과 동일한 구조 (documentOverview, summary, clauseReviews 등)
    긴/구조적 검토 요청 입력에 사용. 일반 짧은 질문은 /v1/chat 사용.
    """
    # TODO: 사용량 제한 임시 비활성화 — 평가 완료 후 복원할 것
    # _check_daily_limit()
    #
    # if user:
    #     user_id = user["sub"]
    #     current = repository.get_usage_count(user_id=user_id)
    #     if current >= USER_LIMIT:
    #         raise HTTPException(status_code=429, detail="LIMIT_EXCEEDED")
    # else:
    #     session_id = x_session_id or "anonymous"
    #     session_count = repository.get_usage_count(session_id=session_id)
    #     if session_count >= GUEST_LIMIT:
    #         raise HTTPException(status_code=429, detail="LOGIN_REQUIRED")
    #     client_ip = _client_ip(request)
    #     ip_key = f"ip:{client_ip}"
    #     if repository.get_usage_count(session_id=ip_key) >= IP_DAILY_LIMIT:
    #         raise HTTPException(status_code=429, detail="LOGIN_REQUIRED")

    text = (payload.text or "").strip()
    if len(text) < 100:
        raise HTTPException(status_code=400, detail="심층 검토는 100자 이상의 본문이 필요합니다.")

    try:
        from law_rag_core.retrieval import RetrievalService
        from law_rag_core.ai import GeminiService
        retrieval = RetrievalService()
        gemini = GeminiService()

        question = (payload.question or "").strip() or "본 내용을 검토하고 법적 리스크와 권고사항을 정리해 주세요."
        if payload.previous_summary.strip():
            question = f"{question}\n\n[이전 검토 결과 요약]\n{payload.previous_summary[:2000]}"
        chunks = retrieval.retrieve(question)

        result = gemini.review_document(
            document_text=text,
            question=question,
            chunks=chunks,
        )

        _record_usage(user, x_session_id, request)

        # 통합 히스토리 저장 (텍스트 검토 유형)
        try:
            repository.save_chat_history(
                user_id=user["sub"] if user else "anonymous",
                question=text[:500],
                response={"summary": result.get("summary", ""), "overall_risk": result.get("overall_risk", "n/a")},
                session_id=x_session_id,
                request_type="review_text",
            )
        except Exception:
            pass

        _record_ops_event(
            "review_text",
            channel="web",
            user_id=user["sub"] if user else None,
            session_id=None if user else (x_session_id or "anonymous"),
            metadata={
                "text_length": len(text),
                "overall_risk": result.get("overall_risk", "n/a"),
            },
        )
        return _format_review_result(result)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


# ── Document Draft Generation ────────────────────────────────


class DocumentDraftRequest(BaseModel):
    document_type: str  # 서류 종류 (자유 입력)
    facts: dict[str, str]
    domain: str


@app.post("/v1/document-draft")
def generate_document_draft(
    payload: DocumentDraftRequest,
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """법률 문서 초안 생성 — 수집된 사실관계 기반으로 지정 서류 종류의 초안 작성.
    수신: DocumentDraftRequest(document_type: 서류 종류, facts: 사실관계 dict, domain: 법률 분야)
    반환: Gemini가 생성한 문서 초안 dict
    """
    try:
        from law_rag_core.ai import GeminiService
        gemini = GeminiService()
        result = gemini.generate_document_draft(
            document_type=payload.document_type,
            facts=payload.facts,
            domain=payload.domain,
        )
        _record_ops_event(
            "document_draft",
            channel="web",
            user_id=user["sub"] if user else None,
            session_id=None if user else (x_session_id or "anonymous"),
            metadata={
                "document_type": payload.document_type,
                "domain": payload.domain,
            },
        )
        return result
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/v1/precedents/sources", response_model=list[PrecedentSource])
def list_precedent_sources():
    """판례 출처 목록 조회 — 사용 가능한 판례 데이터 소스 반환."""
    try:
        return precedent_service.list_sources()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/v1/precedents/search", response_model=PrecedentSearchResponse)
def search_precedents(
    q: str = Query(..., min_length=1),
    source: list[str] | None = Query(default=None),
    display: int = Query(default=4, ge=1, le=20),
    body_search: bool = Query(default=False, alias="bodySearch"),
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """판례 검색 — 키워드(q)로 판례 검색. 출처/건수/본문검색 옵션 지원."""
    try:
        result = precedent_service.search(
            q,
            display=display,
            sources=source,
            body_search=body_search,
        )
        _record_ops_event(
            "search_precedents",
            channel="web",
            user_id=user["sub"] if user else None,
            session_id=None if user else (x_session_id or "anonymous"),
            metadata={
                "result_count": len(result.results),
                "source_count": len(source or []),
                "body_search": body_search,
            },
        )
        return result
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/v1/precedents/{source_id}/{case_id}", response_model=PrecedentDetailResponse)
def get_precedent_detail(source_id: str, case_id: str):
    """판례 상세 조회 — 출처(source_id)와 사건번호(case_id)로 판례 전문 반환."""
    try:
        return precedent_service.get_detail(source_id, case_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/v1/laws/search")
def search_laws(
    q: str = Query(..., min_length=1),
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """법령 검색 — 제목/FTS 검색 + 벡터 유사도 검색 결합. 결과 부족 시 AI 쿼리 보정 자동 적용.
    수신: q(검색어)
    반환: {"items": [...], "corrected": 보정된 검색어(있을 경우)}
    """
    try:
        # 제목/FTS 검색과 벡터 기반 검색을 결합
        title_results = repository.search_laws(q)
        seen_law_ids = {r["law_id"] for r in title_results}

        # 벡터 유사도로 추가 법령 검색
        try:
            from law_rag_core.retrieval import RetrievalService
            retrieval = RetrievalService()
            chunks = retrieval.retrieve(q)
            for chunk in chunks:
                if chunk.law_id not in seen_law_ids:
                    seen_law_ids.add(chunk.law_id)
                    # Fetch law document info for this chunk
                    law_detail = repository.get_law_detail(chunk.law_id)
                    if law_detail and law_detail.get("document"):
                        doc = law_detail["document"]
                        title_results.append({
                            "id": doc.get("id", chunk.id),
                            "law_id": chunk.law_id,
                            "title": chunk.law_title,
                            "law_type": doc.get("law_type"),
                            "status": doc.get("status"),
                            "effective_date": doc.get("effective_date"),
                            "score": round(chunk.score, 2),
                        })
        except Exception as exc:
            import logging
            logging.getLogger(__name__).warning("Vector search in law search failed: %s", exc)

        # 결과 부족 시 AI로 검색어 자동 보정
        corrected = None
        if len(title_results) < 3:
            try:
                from law_rag_core.ai import GeminiService
                gemini = GeminiService()
                corrected = gemini.correct_query(q)
                if corrected:
                    extra = repository.search_laws(corrected)
                    for r in extra:
                        if r["law_id"] not in seen_law_ids:
                            seen_law_ids.add(r["law_id"])
                            title_results.append(r)
            except Exception:
                pass

        result = {"items": title_results}
        if corrected:
            result["corrected"] = corrected
        _record_ops_event(
            "search_laws",
            channel="web",
            user_id=user["sub"] if user else None,
            session_id=None if user else (x_session_id or "anonymous"),
            metadata={
                "result_count": len(title_results),
                "corrected": bool(corrected),
            },
        )
        return result
    except ArtifactNotReadyError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


@app.get("/v1/laws/{law_id}")
def get_law(law_id: str):
    """법령 상세 조회 — law_id로 법령 문서 전체 반환. 없으면 404."""
    try:
        law = repository.get_law_detail(law_id)
    except ArtifactNotReadyError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    if law is None:
        raise HTTPException(status_code=404, detail="Law not found.")
    return law


@app.get("/v1/laws/{law_id}/articles/{article_no}")
def get_article(law_id: str, article_no: str):
    """조문 상세 조회 — 특정 법령의 특정 조문(article_no) 반환. 없으면 404."""
    try:
        article = repository.get_article(law_id, article_no)
    except ArtifactNotReadyError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    if article is None:
        raise HTTPException(status_code=404, detail="Article not found.")
    return article


@app.post("/v1/saved-answers")
def save_answer(payload: SavedAnswerCreate, user: dict | None = Depends(get_optional_user)):
    """답변 저장 — 사용자가 유용한 답변을 북마크/저장."""
    if user and not payload.user_id:
        payload.user_id = user["sub"]
    return repository.create_saved_answer(payload)


@app.get("/v1/saved-answers")
def list_saved_answers(
    user_id: str | None = Query(default=None, alias="userId"),
    user: dict | None = Depends(get_optional_user),
):
    """저장된 답변 목록 조회 — 사용자별 저장 답변 반환."""
    effective_user_id = user_id or (user["sub"] if user else None)
    return {"items": repository.list_saved_answers(user_id=effective_user_id)}


@app.post("/v1/feedback")
def create_feedback(
    payload: FeedbackCreate,
    user: dict | None = Depends(get_optional_user),
    x_session_id: str | None = Header(default=None),
):
    """피드백 생성 — 사용자 피드백(좋아요/싫어요/코멘트) 저장."""
    result = repository.create_feedback(payload)
    _record_ops_event(
        "feedback",
        channel="web",
        user_id=user["sub"] if user else None,
        session_id=None if user else (x_session_id or "anonymous"),
        metadata={"rating": payload.rating},
    )
    return result


# ── Chat History & Thread API ──────────────────────────────


@app.get("/v1/history")
def get_history(
    request_type: str | None = Query(default=None),
    limit: int = Query(default=50, ge=1, le=200),
    user_id: str | None = Query(default=None, alias="userId"),
    user: dict | None = Depends(get_optional_user),
    x_admin_api_key: str | None = Header(default=None),
):
    """대화 히스토리 조회. 관리자 키로 접근 시 전체, 일반 사용자는 본인만."""
    is_admin = settings.admin_api_key and x_admin_api_key == settings.admin_api_key
    if is_admin:
        target_user = user_id or "anonymous"
        # 관리자는 모든 사용자 히스토리 조회 가능 — user_id 미지정 시 최근 전체
        with get_state_db_connection() as conn:
            cursor = conn.cursor()
            if request_type:
                # "chat"이면 chat/deep_dive 모두, "review"이면 review_file/review_text 모두
                if request_type == "chat":
                    type_filter = "request_type IN ('chat', 'deep_dive_follow_up', 'deep_dive_analysis')"
                elif request_type == "review":
                    type_filter = "request_type IN ('review_file', 'review_text')"
                else:
                    type_filter = "request_type = ?"
                if "?" in type_filter:
                    cursor.execute(
                        f"SELECT id, thread_id, turn_index, request_type, question, response, has_attachments, attachment_names, created_at, user_id FROM chat_history WHERE {type_filter} ORDER BY created_at DESC LIMIT ?",
                        (request_type, limit),
                    )
                else:
                    cursor.execute(
                        f"SELECT id, thread_id, turn_index, request_type, question, response, has_attachments, attachment_names, created_at, user_id FROM chat_history WHERE {type_filter} ORDER BY created_at DESC LIMIT ?",
                        (limit,),
                )
            else:
                cursor.execute(
                    "SELECT id, thread_id, turn_index, request_type, question, response, has_attachments, attachment_names, created_at, user_id FROM chat_history ORDER BY created_at DESC LIMIT ?",
                    (limit,),
                )
            import json as _json
            items = []
            for r in cursor.fetchall():
                items.append({
                    "id": r[0], "thread_id": r[1], "turn_index": r[2],
                    "request_type": r[3], "question": r[4],
                    "response": _json.loads(r[5]) if r[5] else {},
                    "has_attachments": bool(r[6]), "attachment_names": _json.loads(r[7]) if r[7] else [],
                    "created_at": r[8], "user_id": r[9],
                })
            cursor.close()
        return {"items": items}
    if not user:
        raise HTTPException(status_code=401, detail="로그인이 필요합니다.")
    return {"items": repository.get_chat_history(user["sub"], limit, request_type=request_type)}


@app.get("/v1/history/thread/{thread_id}")
def get_thread(
    thread_id: str,
    user: dict | None = Depends(get_optional_user),
    x_admin_api_key: str | None = Header(default=None),
):
    """스레드 ID로 대화 전체 턴을 순서대로 조회한다. 관리자 키로도 접근 가능."""
    is_admin = settings.admin_api_key and x_admin_api_key == settings.admin_api_key
    items = repository.get_thread_history(thread_id)
    if not items:
        raise HTTPException(status_code=404, detail="스레드를 찾을 수 없습니다.")
    if not is_admin and (not user or items[0].get("user_id") != user["sub"]):
        raise HTTPException(status_code=403, detail="접근 권한이 없습니다.")
    return {"thread_id": thread_id, "turns": items}


@app.post("/v1/admin/ingest/run", dependencies=[Depends(require_admin)])
def run_ingest(payload: IngestRunRequest):
    """[관리자] 법령 데이터 수집 실행 — 지정 경로의 법령 파일을 파싱/인덱싱."""
    try:
        return ingest_service.ingest_path(
            payload.input_path,
            mode=payload.mode,
            apply_schema=payload.apply_schema,
            reindex=payload.reindex,
        )
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.get("/v1/admin/ingest/jobs/{job_id}", dependencies=[Depends(require_admin)])
def get_ingest_job(job_id: str):
    """[관리자] 수집 작업 상태 조회 — job_id로 진행 상황 확인."""
    try:
        return ingest_service.get_job(job_id)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.get("/v1/admin/stats", dependencies=[Depends(require_admin)])
def admin_stats():
    """[관리자] 시스템 통계 — 문서/청크/저장답변/피드백 수 + 스토리지 용량."""
    try:
        return repository.get_admin_stats()
    except ArtifactNotReadyError:
        return {
            "documents": 0,
            "chunks": 0,
            "saved_answers": 0,
            "feedback": 0,
            "estimated_text_bytes": 0,
            "estimated_vector_bytes": 0,
            "relation_bytes": {
                "laws.sqlite": 0,
                "state.sqlite": settings.state_db_path.stat().st_size if settings.state_db_path.exists() else 0,
                "article_embeddings.npy": 0,
                "article_embedding_ids.json": 0,
                "manifest.json": 0,
            },
        }


@app.get("/v1/admin/dashboard", dependencies=[Depends(require_admin)])
def admin_dashboard():
    """[관리자] 대시보드 — 사용자/질문/캐시/스토리지 통계 + 7일 트렌드 + 최근 질문 목록."""
    snapshot = repository.get_admin_dashboard_snapshot()
    snapshot["policy_status"] = {
        "chat_limit": {
            "label": "채팅 한도",
            "status": "disabled",
            "detail": "/v1/chat 질문 횟수 제한은 현재 코드에서 주석 처리되어 비활성화돼 있습니다.",
        },
        "deep_dive_limit": {
            "label": "심층 질문 한도",
            "status": "disabled",
            "detail": "/v1/deep-dive 질문 횟수 제한은 현재 코드에서 주석 처리되어 비활성화돼 있습니다.",
        },
        "analyze_text_limit": {
            "label": "텍스트 검토 한도",
            "status": "enabled",
            "detail": "/v1/analyze-text는 한도 체크와 usage_counts 누적이 모두 적용됩니다.",
        },
        "analyze_file_limit": {
            "label": "파일 검토 한도",
            "status": "partial",
            "detail": "/v1/analyze-file는 글로벌 일일 한도만 체크하고 usage_counts에는 직접 누적하지 않습니다.",
        },
        "response_cache_write": {
            "label": "응답 캐시 쓰기",
            "status": "disabled",
            "detail": "response_cache는 현재 채팅 파이프라인에서 읽기만 남아 있고 신규 write는 비활성화돼 있습니다.",
        },
    }
    return snapshot


def _chunk_text(text: str, *, width: int = 72) -> list[str]:
    """텍스트를 SSE 스트리밍용 청크로 분할 — 단어 단위로 width(기본 72자) 이하씩 나눔.
    수신: text(원본 텍스트), width(청크 최대 너비)
    반환: 분할된 텍스트 청크 리스트
    """
    words = text.split()
    if not words:
        return []
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for word in words:
        next_len = current_len + len(word) + (1 if current else 0)
        if current and next_len > width:
            chunks.append(" ".join(current))
            current = [word]
            current_len = len(word)
            continue
        current.append(word)
        current_len = next_len
    if current:
        chunks.append(" ".join(current))
    return chunks
