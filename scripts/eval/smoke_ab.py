"""최근 5케이스 스모크 A/B 비교.

A: /tmp/kr-law-rag-A (수정 전 HEAD) 의 review_document 결과
B: 현재 워킹트리 (프롬프트·judgment 수정 후)

- 같은 DB·같은 임베딩 사용 (상대경로 유지)
- 절단 없이 풀텍스트 저장 (변호사 첨부, 시스템 응답 모두)
- 결과물: review-reports/smoke-ab-<timestamp>.md  (4단 비교: 질문 / 변호사답변 / A / B)
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
WORKTREE_A = Path("/tmp/kr-law-rag-A")

TARGETS = [
    "1775637735.186559",
    "1774507045.148539",
    "1767574880.186339",
    "1765778055.035789",
    "1765776630.938769",
]


def load_goldset() -> tuple[dict[str, dict], dict[str, dict]]:
    Q = json.loads((ROOT / ".local-data/eval-questions.json").read_text())
    A = json.loads((ROOT / ".local-data/eval-answers.json").read_text())
    return {q["thread_ts"]: q for q in Q}, {a["thread_ts"]: a for a in A}


def extract_file(path: Path) -> str:
    """절단 없이 전체 텍스트."""
    s = path.suffix.lower()
    try:
        if s == ".pdf":
            import pypdf
            r = pypdf.PdfReader(str(path))
            return "\n".join((p.extract_text() or "") for p in r.pages)
        if s == ".docx":
            from docx import Document
            d = Document(str(path))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        if s in (".doc",):
            # best-effort with libreoffice or antiword
            for cmd in (["antiword", str(path)], ["libreoffice", "--headless", "--convert-to", "txt", "--outdir", "/tmp", str(path)]):
                try:
                    r = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                    if r.returncode == 0 and r.stdout.strip():
                        return r.stdout
                except FileNotFoundError:
                    pass
            return f"(.doc 파싱 실패: {path.name})"
        if s in (".hwp", ".hwpx"):
            try:
                import olefile, zlib
                if not olefile.isOleFile(str(path)):
                    return f"(.hwp OLE 형식 아님: {path.name})"
                ole = olefile.OleFileIO(str(path))
                texts = []
                for i in range(256):
                    nm = f"BodyText/Section{i}"
                    if not ole.exists(nm):
                        break
                    data = ole.openstream(nm).read()
                    try:
                        data = zlib.decompress(data, -15)
                    except zlib.error:
                        pass
                    buf = ""
                    for j in range(0, len(data) - 1, 2):
                        ch = int.from_bytes(data[j:j + 2], "little")
                        if (32 <= ch < 0xFFFF) or ch in (10, 13):
                            buf += chr(ch)
                    texts.append(buf)
                ole.close()
                return "\n".join(texts)
            except Exception as e:
                return f"(.hwp 파싱 실패: {e})"
        return ""
    except Exception as e:
        return f"(파싱 오류: {e})"


def collect_lawyer_fulltext(ans: dict) -> tuple[str, list[str]]:
    body_parts: list[str] = []
    attachment_names: list[str] = []
    for la in ans.get("lawyer_answers") or []:
        if la.get("text"):
            body_parts.append(f"[Slack 본문 — {la.get('author','')}]\n{la['text']}")
        for f in la.get("files") or la.get("attachments") or []:
            if f.get("downloaded") and f.get("local_path"):
                fp = ROOT / f["local_path"]
                if fp.exists():
                    txt = extract_file(fp)
                    attachment_names.append(fp.name)
                    body_parts.append(f"\n=== 첨부: {fp.name} ({len(txt)}자) ===\n{txt}")
    return "\n\n".join(body_parts), attachment_names


def run_system(worktree_root: Path, question_text: str, question_files: list[dict], thread_ts: str) -> dict:
    """해당 worktree의 Python으로 review_document 호출."""
    script = """
import sys, json, os
sys.path.insert(0, os.path.join(os.environ['WKTREE'], 'packages', 'python', 'src'))
# Working dir을 worktree로 바꿔서 상대경로 설정 로드되게
os.chdir(os.environ['WKTREE'])

from law_rag_core.retrieval.service import RetrievalService
from law_rag_core.ai import GeminiService

q_text = sys.stdin.read()
payload = json.loads(q_text)
question = payload["question"]
files = payload.get("files") or []
rs = RetrievalService()
gs = GeminiService()
chunks = rs.retrieve(question[:500] if question else "")
if files:
    doc_list = [{"filename": f["name"], "text": f["text"]} for f in files]
    if len(doc_list) == 1:
        r = gs.review_document(document_text=doc_list[0]["text"], question=question, chunks=chunks)
    else:
        r = gs.review_documents_orchestrated(documents=doc_list, question=question, chunks=chunks)
else:
    # 질문만 텍스트로 → review_document(document_text=question)로 돌림
    r = gs.review_document(document_text=question, question=question, chunks=chunks)
print(json.dumps(r, ensure_ascii=False))
"""
    # 질문 첨부 파일 텍스트 추출
    files_payload = []
    for f in question_files:
        if f.get("downloaded") and f.get("local_path"):
            fp = ROOT / f["local_path"]
            if fp.exists():
                files_payload.append({"name": f["name"], "text": extract_file(fp)})

    payload = {"question": question_text, "files": files_payload}
    env = os.environ.copy()
    env["WKTREE"] = str(worktree_root)
    env["PYTHONPATH"] = str(worktree_root / "packages" / "python" / "src")
    # .venv는 현재 프로젝트 것 공유 (worktree는 코드만 다름, 패키지는 동일)
    py = str(ROOT / ".venv" / "bin" / "python")
    try:
        r = subprocess.run(
            [py, "-c", script],
            input=json.dumps(payload),
            capture_output=True, text=True, timeout=900, env=env,
        )
        if r.returncode != 0:
            return {"_error": f"returncode={r.returncode}", "_stderr": r.stderr[-2000:]}
        # stdout 마지막 JSON 한 줄만
        out = r.stdout.strip().split("\n")[-1]
        return json.loads(out)
    except Exception as e:
        return {"_error": str(e)}


def serialize_result(r: dict) -> str:
    """review_document 결과를 사람이 읽을 수 있는 풀텍스트로 직렬화 (절단 없음)."""
    if "_error" in r:
        return f"[실행 오류] {r['_error']}\n{r.get('_stderr','')}"
    out: list[str] = []
    ov = r.get("document_overview") or {}
    if ov:
        out.append(f"**문서 성격**: {ov.get('nature','')}")
        if ov.get("principals"):
            out.append(f"**주요 주체**: {', '.join(ov['principals'])}")
        if ov.get("standpoint"):
            out.append(f"**검토 관점**: {ov['standpoint']}")
    axes = r.get("review_axes") or []
    if axes:
        out.append(f"**검토 축**: {' / '.join(axes)}")
    out.append(f"**전체 위험도**: {r.get('overall_risk','?')}")
    if r.get("summary"):
        out.append(f"\n### 요약\n{r['summary']}")

    frame = r.get("institutional_frame")
    if frame:
        out.append(f"\n### 제도 프레임\n{frame}")
    inst_axes = r.get("axes") or []
    if inst_axes:
        out.append("\n### 선언된 축")
        for ax in inst_axes:
            out.append(f"- **{ax.get('name','?')}** — {ax.get('why','')}")

    # judgment (다중)
    j = r.get("judgment") or {}
    if j.get("is_judgment_question"):
        out.append("\n### 판단")
        jarr = j.get("judgments")
        if not jarr and j.get("verdict"):
            jarr = [j]
        for i, jj in enumerate(jarr or [], 1):
            label = jj.get("issue_label") or f"판단 {i}"
            out.append(f"\n**{label}** [{jj.get('verdict','?')}]")
            if jj.get("short_answer"):
                out.append(f"- {jj['short_answer']}")
            if jj.get("reasoning"):
                out.append(f"- 근거: {jj['reasoning']}")
            for a in jj.get("key_authorities", []) or []:
                kind = "판례" if a.get("type") == "case" else "법령"
                out.append(f"  - [{kind}] {a.get('ref','?')} — {a.get('point','')}")

    rej = r.get("rejected_citations") or []
    if rej:
        out.append("\n### 적용되지 않는 인용")
        for x in rej:
            out.append(f"- {x.get('ref','?')}: {x.get('reason','')}")

    obs = r.get("observations") or r.get("clause_reviews") or []
    if obs:
        out.append(f"\n### 상세 검토 ({len(obs)}건)")
        for i, o in enumerate(obs, 1):
            sev = o.get("severity") or o.get("risk_level", "?")
            loc = o.get("locator") or (f"제{o.get('clause_no')}조" if o.get("clause_no") else "?")
            title = o.get("locator_title") or o.get("clause_title") or ""
            out.append(f"\n**{i}. [{sev}] {loc} {title}**")
            if o.get("original_text"):
                out.append(f"원문: \"{o['original_text']}\"")
            issues = o.get("issues") or []
            if not issues and (o.get("concern") or o.get("issue") or o.get("suggestion")):
                issues = [{"severity": sev, "concern": o.get("concern") or o.get("issue"),
                           "suggestion": o.get("suggestion"), "basis": o.get("basis") or o.get("legal_basis")}]
            for it in issues:
                concern = it.get("concern") or it.get("issue") or ""
                out.append(f"- 문제: {concern}")
                if it.get("suggestion"):
                    out.append(f"- 제안: {it['suggestion']}")
                basis = it.get("basis") or it.get("legal_basis")
                if basis:
                    out.append(f"- 근거: {basis}")

    gaps = r.get("gaps") or []
    if gaps:
        out.append(f"\n### 부족·보완 ({len(gaps)}건)")
        for g in gaps:
            out.append(f"- **{g.get('topic','?')}**: {g.get('reason','')}")
            if g.get("suggestion"):
                out.append(f"  - 제안: {g['suggestion']}")

    ext = r.get("external_considerations") or []
    if ext:
        out.append(f"\n### 문서 외 고려사항 ({len(ext)}건)")
        for e in ext:
            title = e.get("section_title") or ""
            out.append(f"- **[{title}] {e.get('topic','?')}**: {e.get('detail','')}")

    scen = r.get("risk_scenarios") or []
    if scen:
        out.append(f"\n### 리스크 시나리오 ({len(scen)}건)")
        for s in scen:
            out.append(f"- [{s.get('severity','?')}] {s.get('trigger','?')}")
            if s.get("root_cause"):
                out.append(f"  - 원인: {s['root_cause']}")
            if s.get("suggestion"):
                out.append(f"  - 예방: {s['suggestion']}")

    if r.get("action_items"):
        out.append("\n### 행동 항목")
        for a in r["action_items"]:
            out.append(f"- {a}")

    return "\n".join(out)


def main() -> None:
    questions, answers = load_goldset()
    ts_stamp = datetime.now().strftime("%Y-%m-%dT%H-%M-%S")
    report_path = ROOT / "review-reports" / f"smoke-ab-{ts_stamp}.md"
    report_path.parent.mkdir(exist_ok=True)

    lines: list[str] = []
    lines.append(f"# 스모크 A/B 비교 — {ts_stamp}")
    lines.append("")
    lines.append("- **A**: `/tmp/kr-law-rag-A` (수정 전 HEAD `1ebb552`)")
    lines.append("- **B**: 현재 워킹트리 (프롬프트 3개 + gemini.py review_judgment 배열화 + UI 4개 수정)")
    lines.append("- 대상: 최근 5케이스 (thread_ts 내림차순)")
    lines.append("")
    lines.append("---")

    for i, ts in enumerate(TARGETS, 1):
        print(f"\n[{i}/{len(TARGETS)}] {ts}", flush=True)
        q = questions.get(ts) or {}
        a = answers.get(ts) or {}
        qobj = q.get("question") or {}
        q_text = qobj.get("text") or ""
        q_files = qobj.get("files") or []
        author = qobj.get("author") or "?"
        posted = q.get("posted_at") or ""

        lawyer_text, att_names = collect_lawyer_fulltext(a)

        print(f"  → A run (worktree) ...", flush=True)
        t0 = time.time()
        a_result = run_system(WORKTREE_A, q_text, q_files, ts)
        a_elapsed = int(time.time() - t0)

        print(f"  → B run (current) ...", flush=True)
        t0 = time.time()
        b_result = run_system(ROOT, q_text, q_files, ts)
        b_elapsed = int(time.time() - t0)

        lines.append(f"\n## 케이스 {i}/{len(TARGETS)} — `{ts}`")
        lines.append(f"- 작성자: {author} / 게시: {posted}")
        lines.append(f"- 질문 첨부: {len(q_files)}건")
        lines.append(f"- 변호사 답변 첨부: {', '.join(att_names) or '(없음)'}")
        lines.append(f"- A 실행: {a_elapsed}초 / B 실행: {b_elapsed}초")

        lines.append("\n<details><summary>📩 질문 (풀텍스트)</summary>\n")
        lines.append(f"```\n{q_text}\n```")
        lines.append("\n</details>\n")

        lines.append("\n<details><summary>⚖️ 변호사 답변 (Slack 본문 + 첨부 풀텍스트)</summary>\n")
        lines.append(f"```\n{lawyer_text}\n```")
        lines.append("\n</details>\n")

        lines.append("\n<details><summary>🅰️ A 응답 (수정 전)</summary>\n")
        lines.append(serialize_result(a_result))
        lines.append("\n</details>\n")

        lines.append("\n<details><summary>🅱️ B 응답 (수정 후)</summary>\n")
        lines.append(serialize_result(b_result))
        lines.append("\n</details>\n")

        lines.append("\n---")

        # 중간 저장 (긴 실행 중 중단돼도 부분 결과 보존)
        report_path.write_text("\n".join(lines))

    report_path.write_text("\n".join(lines))
    print(f"\n✅ 리포트 저장: {report_path}")
    print(f"   크기: {report_path.stat().st_size // 1024}KB")


if __name__ == "__main__":
    main()
