"""
터미널에서 바로 질문하고 결과 확인하는 스크립트.

사용법:
  .venv/bin/python scripts/ask.py "법률 질문을 입력하세요"
  .venv/bin/python scripts/ask.py --file 계약서.pdf "이 계약서에서 문제될 부분은?"
  .venv/bin/python scripts/ask.py --file a.docx --file b.docx "두 문서 비교해줘"
  .venv/bin/python scripts/ask.py   ← 대화형 모드
"""

from __future__ import annotations

import sys
import time
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "packages" / "python" / "src"))

from law_rag_core.retrieval.service import RetrievalService
from law_rag_core.ai import GeminiService


def extract_text_from_file(file_path: str) -> str:
    """파일에서 텍스트 추출. txt, pdf, docx, 이미지 지원."""
    p = Path(file_path).resolve()
    if not p.exists():
        print(f"  파일 없음: {p}")
        return ""

    suffix = p.suffix.lower()
    size_mb = p.stat().st_size / 1024 / 1024

    # 텍스트 파일
    if suffix in (".txt", ".md", ".csv", ".json"):
        text = p.read_text(encoding="utf-8", errors="ignore")
        print(f"  텍스트 로드: {p.name} ({size_mb:.1f}MB, {len(text)}자)")
        return text

    # HWP 문서
    if suffix in (".hwp", ".hwpx"):
        try:
            import olefile
            import zlib
            if not olefile.isOleFile(str(p)):
                print(f"  HWP 파일 형식이 올바르지 않습니다: {p.name}")
                return ""
            ole = olefile.OleFileIO(str(p))
            text_parts = []
            for i in range(256):
                stream_name = f"BodyText/Section{i}"
                if not ole.exists(stream_name):
                    break
                data = ole.openstream(stream_name).read()
                try:
                    data = zlib.decompress(data, -15)
                except zlib.error:
                    pass
                text = ""
                pos = 0
                while pos < len(data):
                    if pos + 1 < len(data):
                        ch = int.from_bytes(data[pos:pos+2], 'little')
                        if (32 <= ch < 0xFFFF) or ch in (10, 13):
                            text += chr(ch)
                    pos += 2
                text_parts.append(text)
            ole.close()
            text = "\n".join(text_parts)
            text = text.encode("utf-8", errors="ignore").decode("utf-8")
            text = "".join(ch for ch in text if ord(ch) >= 32 or ch in "\n\r\t")
            if text.strip():
                print(f"  HWP 로드: {p.name} ({len(text)}자)")
                return text
            print(f"  HWP 텍스트 추출 실패: {p.name}")
            return ""
        except ImportError:
            print("  olefile 미설치. pip install olefile 필요")
            return ""

    # Word 문서 (.docx)
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(p))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs)
            print(f"  DOCX 로드: {p.name} ({len(doc.paragraphs)}단락, {len(doc.tables)}표, {len(text)}자)")
            return text
        except ImportError:
            print("  python-docx 미설치. pip install python-docx 필요")
            return ""

    # Word 문서 (.doc 구형식)
    if suffix == ".doc":
        import subprocess
        # antiword 시도
        try:
            result = subprocess.run(["antiword", str(p)], capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                print(f"  DOC 로드 (antiword): {p.name} ({len(text)}자)")
                return text
        except FileNotFoundError:
            pass
        # libreoffice 변환 시도
        try:
            import tempfile, os
            out_dir = tempfile.mkdtemp()
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", out_dir, str(p)],
                capture_output=True, timeout=30,
            )
            txt_path = os.path.join(out_dir, p.stem + ".txt")
            if os.path.exists(txt_path):
                with open(txt_path, encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                print(f"  DOC 로드 (libreoffice): {p.name} ({len(text)}자)")
                return text
        except FileNotFoundError:
            pass
        print(f"  .doc 파일을 처리할 수 없습니다. .docx로 변환 후 다시 시도하세요.")
        print(f"  (antiword 또는 libreoffice 필요: brew install antiword)")
        return ""

    # PDF
    if suffix == ".pdf":
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(str(p)) as pdf:
                for page in pdf.pages[:30]:
                    text_parts.append(page.extract_text() or "")
            text = "\n".join(text_parts)
            print(f"  PDF 로드: {p.name} ({len(pdf.pages)}페이지, {len(text)}자)")
            return text
        except ImportError:
            print("  pdfplumber 미설치. pip install pdfplumber 필요")
            return ""

    # 이미지 (Gemini Vision OCR)
    if suffix in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        gemini = GeminiService()
        if not gemini.client:
            print("  Gemini API 키 없음")
            return ""
        try:
            from google.genai import types
            with open(str(p), "rb") as f:
                image_bytes = f.read()
            mime = f"image/{suffix[1:]}" if suffix != ".jpg" else "image/jpeg"
            response = gemini.client.models.generate_content(
                model=gemini.settings.gemini_model,
                contents=[
                    types.Part.from_bytes(data=image_bytes, mime_type=mime),
                    "이 이미지의 텍스트를 모두 추출하라. 표는 표 형태로, 도장/서명은 [도장]/[서명]으로 표시. 텍스트만 반환.",
                ],
                config=types.GenerateContentConfig(temperature=0.0),
            )
            text = response.text or ""
            print(f"  이미지 OCR: {p.name} ({size_mb:.1f}MB → {len(text)}자)")
            return text
        except Exception as e:
            print(f"  이미지 OCR 실패: {e}")
            return ""

    print(f"  지원하지 않는 형식: {suffix}")
    return ""


def ask(question: str, file_texts: list[tuple[str, str]] | None = None, previous_review: dict | None = None) -> dict | None:
    """질문 + 첨부 파일들로 RAG 답변 생성.

    file_texts: [(파일명, 추출텍스트), ...]
    파일 첨부 시 → 문서 검토 모드 (조항별 검토)
    파일 없을 시 → 일반 질문 모드 (법률 Q&A)
    """
    rs = RetrievalService()
    gemini = GeminiService()

    file_texts = file_texts or []
    has_files = bool(file_texts)

    combined_file_text = ""
    if file_texts:
        parts = []
        for fname, ftext in file_texts:
            parts.append(f"--- {fname} ---\n{ftext}")
        combined_file_text = "\n\n".join(parts)

    print(f"\n{'='*60}")
    print(f"  질문: {question[:80]}{'...' if len(question) > 80 else ''}")
    if file_texts:
        for fname, ftext in file_texts:
            print(f"  첨부: {fname} ({len(ftext)}자)")
        print(f"  모드: 문서 검토 (조항별 분석)")
    else:
        print(f"  모드: 일반 질문")
    print(f"{'='*60}")

    # RAG 검색
    search_query = question or combined_file_text[:200]
    start = time.time()
    chunks = rs.retrieve(search_query)
    search_ms = int((time.time() - start) * 1000)

    print(f"\n검색 결과 ({len(chunks)}건, {search_ms}ms)")
    for c in chunks:
        print(f"  [{c.law_title}] {c.article_no or ''} {c.article_title or ''} (score: {c.score:.1f})")

    start = time.time()

    if has_files:
        # ── 문서 검토 모드 ──
        q = question or "이 문서를 법적으로 검토해 주세요."
        if len(file_texts) == 1:
            result = gemini.review_document(
                document_text=combined_file_text,
                question=q,
                chunks=chunks,
                previous_review=previous_review,
            )
        else:
            doc_list = [{"filename": fname, "text": ftext} for fname, ftext in file_texts]
            result = gemini.review_documents_orchestrated(
                documents=doc_list,
                question=q,
                chunks=chunks,
            )
        llm_ms = int((time.time() - start) * 1000)

        overview = result.get("document_overview") or {}
        nature = overview.get("nature") or result.get("document_type", "unknown")
        overall_risk = result.get("overall_risk", "unknown")
        risk_tag_map = {"high": "HIGH", "medium": "MEDIUM", "low": "LOW", "ok": "OK", "n/a": "N/A"}

        print(f"\n문서 성격: {nature}")
        principals = overview.get("principals") or result.get("parties", []) or []
        if principals:
            print(f"주요 주체: {' / '.join(principals)}")
        standpoint = overview.get("standpoint") or result.get("review_perspective", "")
        if standpoint:
            print(f"검토 관점: {standpoint}")
        axes = result.get("review_axes") or []
        if axes:
            print(f"검토 축: {' · '.join(axes)}")
        print(f"전체 주의 정도: {risk_tag_map.get(overall_risk, overall_risk)}")

        detected = result.get("detected_clauses", [])
        if detected:
            print(f"발견 단위: {len(detected)}개 — {', '.join(detected[:20])}")

        judgment = result.get("judgment") or {}
        if judgment.get("is_judgment_question"):
            verdict_label_map = {
                "likely_yes": "가능성 높음",
                "likely_no": "해당 가능성 낮음",
                "depends": "조건부",
                "needs_more_info": "정보 부족",
                "not_applicable": "해당 없음",
            }
            judgments_arr = judgment.get("judgments") or []
            # 구 응답 호환: 배열이 비면 최상위 필드로 단일 판단 구성
            if not judgments_arr and judgment.get("verdict"):
                judgments_arr = [{
                    "issue_label": "",
                    "verdict": judgment.get("verdict"),
                    "short_answer": judgment.get("short_answer"),
                    "reasoning": judgment.get("reasoning"),
                    "key_authorities": judgment.get("key_authorities") or [],
                    "missing_facts": judgment.get("missing_facts") or [],
                    "typical_path": judgment.get("typical_path") or "",
                }]
            for idx, j in enumerate(judgments_arr, 1):
                verdict = j.get("verdict", "?")
                label = verdict_label_map.get(verdict, verdict)
                header = f"\n[결론"
                if len(judgments_arr) > 1:
                    tag = j.get("issue_label") or f"판단 {idx}"
                    header += f" — {tag}"
                header += f": {label}]"
                print(header)
                short = j.get("short_answer") or ""
                if short:
                    print(f"  {short}")
                if j.get("reasoning"):
                    print(f"  근거: {j['reasoning']}")
                for a in j.get("key_authorities", []) or []:
                    kind = "판례" if a.get("type") == "case" else "법령"
                    print(f"    [{kind}] {a.get('ref', '?')} — {a.get('point', '')}")
                mf = j.get("missing_facts") or []
                if mf:
                    print("  필요한 추가 사실:")
                    for f in mf:
                        print(f"    - {f}")
                if j.get("typical_path"):
                    print(f"  일반 경로: {j['typical_path']}")
        rejected = (result.get("rejected_citations") or []) + (judgment.get("rejected_citations") or [])
        if rejected:
            seen_refs = set()
            printed = False
            for r in rejected:
                ref = r.get("ref") or ""
                if not ref or ref in seen_refs:
                    continue
                seen_refs.add(ref)
                if not printed:
                    print("\n[적용되지 않는 인용 법령]")
                    printed = True
                print(f"  - {ref}: {r.get('reason', '')}")

        frame = result.get("institutional_frame") or ""
        if frame:
            print(f"\n제도 프레임\n  {frame}")
        inst_axes = result.get("axes") or []
        if inst_axes:
            for ax in inst_axes:
                print(f"  - {ax.get('name', '?')}: {ax.get('why', '')}")

        print(f"\n요약")
        print(f"  {result.get('summary', '')}")

        observations = result.get("observations") or result.get("clause_reviews") or []
        if observations:
            print(f"\n상세 검토 ({len(observations)}건)")
            for i, o in enumerate(observations, 1):
                sev = o.get("severity") or o.get("risk_level", "?")
                tag = {"high": "[HIGH]", "medium": "[MED]", "low": "[LOW]", "ok": "[OK]"}.get(sev, f"[{sev}]")
                locator = o.get("locator") or (f"제{o.get('clause_no')}조" if o.get("clause_no") is not None else "?")
                title = o.get("locator_title") or o.get("clause_title") or ""
                print(f"\n  {i}. {tag} {locator} {title}")
                if o.get("original_text"):
                    ot = o["original_text"]
                    print(f"     원문: \"{ot[:120]}{'...' if len(ot) > 120 else ''}\"")
                issues = o.get("issues") or []
                if not issues and (o.get("concern") or o.get("issue") or o.get("suggestion")):
                    issues = [{
                        "severity": sev,
                        "concern": o.get("concern") or o.get("issue"),
                        "suggestion": o.get("suggestion"),
                        "basis": o.get("basis") or o.get("legal_basis"),
                    }]
                for j, it in enumerate(issues, 1):
                    isev = it.get("severity", "?")
                    sev_tag = {"high": "[HIGH]", "medium": "[MED]", "low": "[LOW]"}.get(isev, f"[{isev}]")
                    prefix = f"     {j}) {sev_tag}" if len(issues) > 1 else "    "
                    concern = it.get("concern") or it.get("issue")
                    basis = it.get("basis") or it.get("legal_basis")
                    if concern:
                        print(f"{prefix} 문제: {concern}")
                    if it.get("suggestion"):
                        print(f"     {'   ' if len(issues) > 1 else ''}제안: {it['suggestion']}")
                    if basis:
                        print(f"     {'   ' if len(issues) > 1 else ''}근거: {basis}")

        unreviewed = result.get("unreviewed_clauses", [])
        if unreviewed:
            print(f"\n미검토 조항: {', '.join(str(u) for u in unreviewed)}")

        gaps = result.get("gaps") or [
            {"topic": m.get("clause"), "reason": m.get("reason"), "suggestion": m.get("suggestion")}
            for m in (result.get("missing_clauses") or [])
        ]
        if gaps:
            print(f"\n부족·보완 ({len(gaps)}건)")
            for g in gaps:
                print(f"  - {g.get('topic', '?')}: {g.get('reason', '')}")
                if g.get("suggestion"):
                    print(f"    제안: {g['suggestion']}")

        externals = result.get("external_considerations") or []
        if externals:
            print(f"\n문서 외 고려사항 ({len(externals)}건)")
            for ec in externals:
                title = ec.get("section_title") or ""
                label = f"[{title}] " if title else ""
                print(f"  - {label}{ec.get('topic', '?')}: {ec.get('detail', '')}")
                if ec.get("suggestion"):
                    print(f"    권고: {ec['suggestion']}")

        scenarios = result.get("risk_scenarios") or []
        if scenarios:
            print(f"\n예상 리스크 시나리오 ({len(scenarios)}건)")
            for s in scenarios:
                sev = s.get("severity", "?")
                tag = {"high": "[HIGH]", "medium": "[MED]", "low": "[LOW]"}.get(sev, f"[{sev}]")
                print(f"  {tag} {s.get('trigger', '?')}")
                if s.get("root_cause"):
                    print(f"    원인: {s['root_cause']}")
                if s.get("suggestion"):
                    print(f"    예방: {s['suggestion']}")

        # 종합 검토 전용 필드 (오케스트레이션 결과)
        per_doc = result.get("per_document_highlights", [])
        if per_doc:
            print(f"\n파일별 요약 ({len(per_doc)}개)")
            for pd in per_doc:
                print(f"\n  [{pd.get('filename', '?')}] ({pd.get('document_type', '')})")
                for r in pd.get("key_risks", []):
                    print(f"    리스크: {r}")
                for o in pd.get("key_ok", []):
                    print(f"    OK: {o}")

        cross = result.get("cross_document_issues", [])
        if cross:
            print(f"\n문서 간 이슈 ({len(cross)}건)")
            for c in cross:
                docs = ", ".join(c.get("documents_involved", []))
                print(f"  - [{docs}] {c.get('issue', '')}")
                if c.get("detail"):
                    print(f"    {c['detail']}")
                if c.get("suggestion"):
                    print(f"    제안: {c['suggestion']}")

        corrected = result.get("corrected_missing", [])
        if corrected:
            print(f"\n오판 수정 ({len(corrected)}건)")
            for cm in corrected:
                print(f"  - \"{cm.get('originally_claimed_missing', '')}\" → 실제로 {cm.get('actually_found_in', '')}에 있음")

        negotiation = result.get("key_negotiation_points", [])
        if negotiation:
            print(f"\n협상 포인트 (중요도순)")
            for n in negotiation:
                print(f"  - {n}")

        actions = result.get("action_items", [])
        if actions:
            print(f"\n행동 항목")
            for a in actions:
                print(f"  - {a}")

        print(f"\n{result.get('disclaimer', '')}")

    else:
        # ── 일반 질문 모드 ──
        result = gemini.generate_grounded_answer(question=question, chunks=chunks)
        llm_ms = int((time.time() - start) * 1000)

        print(f"\n요약 (grounded={result['grounded']}, {llm_ms}ms)")
        print(f"  {result['summary']}")

        print(f"\n답변")
        for line in result['answer'].split('\n'):
            print(f"  {line}")

        if result.get("decision_factors"):
            print(f"\n판단 요소")
            for f in result["decision_factors"]:
                print(f"  - {f}")

        if result.get("action_items"):
            print(f"\n행동 항목")
            for a in result["action_items"]:
                print(f"  - {a}")

        if result.get("citations"):
            print(f"\n인용 ({len(result['citations'])}건)")
            for c in result["citations"]:
                print(f"  - {c['reason']}")

    print(f"\n총 {search_ms + llm_ms}ms (검색 {search_ms}ms + LLM {llm_ms}ms)")
    print()

    # 문서 검토 결과를 반환 (추가 질문 시 사용)
    if has_files:
        return result
    return None


if __name__ == "__main__":
    import logging
    logging.disable(logging.INFO)

    # 인자 파싱: --file 또는 파일 확장자 자동 감지
    FILE_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".doc", ".hwp", ".hwpx", ".png", ".jpg", ".jpeg", ".webp", ".gif"}

    args = sys.argv[1:]
    file_texts: list[tuple[str, str]] = []
    question_parts: list[str] = []

    i = 0
    while i < len(args):
        if args[i] == "--file" and i + 1 < len(args):
            fp = args[i + 1]
            text = extract_text_from_file(fp)
            if text:
                file_texts.append((Path(fp).name, text))
            i += 2
        elif Path(args[i]).suffix.lower() in FILE_EXTENSIONS and Path(args[i]).exists():
            # 파일 확장자 + 존재하면 자동으로 파일로 인식
            fp = args[i]
            text = extract_text_from_file(fp)
            if text:
                file_texts.append((Path(fp).name, text))
            i += 1
        else:
            question_parts.append(args[i])
            i += 1

    if question_parts or file_texts:
        ask(" ".join(question_parts), file_texts)
    else:
        print("Lemini RAG 테스트")
        print("  종료: q")
        print("  파일 첨부: file: /경로/파일.docx (여러 개는 여러 줄)")
        print("  첨부 완료 후 질문 입력")
        print("  이전 문서 검토 결과 기반 추가 질문 가능")
        print()

        pending_files: list[tuple[str, str]] = []
        last_review: dict | None = None  # 이전 문서 검토 결과
        last_file_texts: list[tuple[str, str]] = []  # 이전 첨부 파일

        while True:
            try:
                if pending_files:
                    prompt = f"질문 ({len(pending_files)}개 첨부됨, 추가: file:, 질문 입력)> "
                elif last_review:
                    prompt = "질문 (이전 검토 기반 추가 질문 가능, 새 file: 로 초기화)> "
                else:
                    prompt = "질문> "
                q = input(prompt).strip()
            except (EOFError, KeyboardInterrupt):
                break
            if not q or q.lower() == "q":
                break

            # 파일 감지: "file:" 접두어 또는 파일 확장자+존재 자동 감지
            detected_file = None
            if q.startswith("file:"):
                detected_file = q[5:].strip()
            else:
                # 입력이 파일 경로처럼 보이면 자동 감지
                candidate = q.strip().strip('"').strip("'")
                if Path(candidate).suffix.lower() in FILE_EXTENSIONS and Path(candidate).exists():
                    detected_file = candidate

            if detected_file:
                text = extract_text_from_file(detected_file)
                if text:
                    pending_files.append((Path(detected_file).name, text))
                    last_review = None
                    print(f"  첨부 완료 ({len(pending_files)}개). 질문을 입력하거나 파일 경로를 추가 입력.")
                continue

            # 질문 실행
            if pending_files:
                # 새 파일 + 질문
                result = ask(q, pending_files)
                last_file_texts = pending_files[:]
                last_review = result
                pending_files = []
            elif last_review and last_file_texts:
                # 이전 결과 기반 추가 질문 (같은 문서에 대해)
                print("  (이전 검토 결과를 참고하여 추가 분석)")
                result = ask(q, last_file_texts, previous_review=last_review)
                last_review = result
            else:
                # 일반 질문
                ask(q)
